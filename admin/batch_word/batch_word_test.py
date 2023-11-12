#!/usr/bin/env python3
# -*- encoding: utf-8 -*-
"""
__file__   : extract_nums_info
__time__   : 2022/8/12 20:51
__author__ : secoder
__resume__ : None
"""
import os
from docx import Document
import xlwt


class ExtractNumsInfo(object):
    def __init__(self, path_name):
        self.path_name = path_name
        self.files = os.listdir(path_name)
        print("当前目录下的文件有：", self.files)

    def extract_info(self):
        doc_files = self.files
        workbook = xlwt.Workbook(encoding='ascii')
        for doc_file in doc_files:
            print("读取的文件名：", doc_file)
            worksheet = workbook.add_sheet(doc_file)
            worksheet.write(0, 0, "监测编号")
            worksheet.write(0, 1, "变化后权属单位（个人）名称")
            worksheet.write(0, 2, "农用地转用、征地和供地批文号")
            worksheet.write(0, 3, "用途说明（建设动工、竣工时间及现状）")
            # 获取文件的当前路径
            path = os.path.join(self.path_name, doc_file)
            print("path" + path)
            document = Document(path)
            tables = document.tables
            print(type(tables))
            print(len(tables))
            for i in range(len(tables)):
                table = tables[i]
                num = table.cell(0, 1).text
                print("num: ", num)
                person = table.cell(1, 1).text
                print("person: ", person)
                document = table.cell(2, 1).text
                print("document: ", document)
                explain = table.cell(3, 1).text
                print("explain: ", explain)
                print("I: ", i)
                worksheet.write(i+1, 0, num)
                worksheet.write(i+1, 1, person)
                worksheet.write(i+1, 2, document)
                worksheet.write(i+1, 3, explain)
            print("读取文件 %s 读取完毕。。。：" % doc_file)    
        workbook.save('./output/2022.xlsx')
        print("文件提取完毕，请查看extract文件夹下的xls文件")
        
    def extract_info2(self):
        doc_files = self.files
        workbook = xlwt.Workbook(encoding='ascii')
        worksheet = workbook.add_sheet("MySheet")
        worksheet.write(0, 0, "监测编号")
        worksheet.write(0, 1, "变化后权属单位（个人）名称")
        worksheet.write(0, 2, "农用地转用、征地和供地批文号")
        worksheet.write(0, 3, "用途说明（建设动工、竣工时间及现状）")
        for doc_file in doc_files:
            # 获取文件的当前路径
            path = os.path.join(self.path_name, doc_file)
            print("path" + path)
            document = Document(path)
            tables = document.tables
            print(type(tables))
            print(len(tables))
            
            for i in range(len(tables)):
                table = tables[i]
                num = table.cell(0, 1).text
                print("num: ", num)
                person = table.cell(1, 1).text
                print("person: ", person)
                document = table.cell(2, 1).text
                print("document: ", document)
                explain = table.cell(3, 1).text
                print("explain: ", explain)
                print("I: ", i)
            print("读取文件 %s 读取完毕。。。：" % doc_file)  
        worksheet.write(i+1, 0, num)
        worksheet.write(i+1, 1, person)
        worksheet.write(i+1, 2, document)
        worksheet.write(i+1, 3, explain)
        workbook.save('./output/2022.xlsx')
        print("文件提取完毕，请查看extract文件夹下的xls文件")


if __name__ == '__main__':
    docs_path_name = r'./docs/'
    eui = ExtractNumsInfo(docs_path_name)
    eui.extract_info2()