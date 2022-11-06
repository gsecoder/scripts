#!/usr/bin/env python3
# -*- encoding: utf-8 -*-
"""
__file__   : extract_nums_info
__time__   : 2022/5/15 20:51
__author__ : secoder
__resume__ : None
"""
import os
import time

from docx import Document
import xlwt
import xlrd
import pandas as pd
# from  pandas import DataFrame


class ExtractNumsInfo(object):
    def __init__(self, path_name):
        self.path_name = path_name
        self.files = os.listdir(path_name)
        self.excel_name = './extract_out/2022.xlsx'
        # 创建新的sheet表

    def extract_info(self):
        doc_files = self.files
        workbook = xlwt.Workbook(encoding='ascii')
        for doc_file in doc_files:
            print("读取的文件名：", doc_file)
            worksheet = workbook.add_sheet(doc_file)
            worksheet.write(0, 0, "监测编号")
            worksheet.write(0, 1, "变化后权属单位（个人）名称")
            worksheet.write(0, 2, "农用地转用、征地和供地批文号")
            worksheet.write(0, 3, "用途说明（建设动工、竣工时间及现状）")
            # 获取文件的当前路径
            path = os.path.join(self.path_name, doc_file)
            print("path" + path)
            document = Document(path)
            tables = document.tables
            print(type(tables))
            print(len(tables))
            for i in range(len(tables)):
                table = tables[i]
                num = table.cell(0, 1).text
                print("num: ", num)
                person = table.cell(1, 1).text
                print("person: ", person)
                document = table.cell(2, 1).text
                print("document: ", document)
                explain = table.cell(3, 1).text
                print("explain: ", explain)
                print("I: ", i)
                worksheet.write(i+1, 0, num)
                worksheet.write(i+1, 1, person)
                worksheet.write(i+1, 2, document)
                worksheet.write(i+1, 3, explain)
            print("读取文件 %s 读取完毕。。。：" % doc_file)
        workbook.save(self.excel_name)
        print("文件提取完毕，请查看extract文件夹下的xls文件")

    def merge_excel_sheet(self):
        # 读入数据文件
        sheets = pd.read_excel(self.excel_name, None)
        keys = list(sheets.keys())
        # 数据合并
        iris_concat = pd.DataFrame()
        for i in keys:
            iris1 = sheets[i]
            iris_concat = pd.concat([iris_concat,iris1])
        iris_concat.to_excel("./extract_out/2022_end.xlsx")
        print("sheet页合并完成，请查看extract_out/目录下的[2022_end.xlsx文件即为提取文件")


    def extract_info2(self):
        doc_files = self.files
        doc_table_nums = 0
        for doc_file in doc_files:
            print("doc_file", doc_file)
            doc_path = os.path.join(self.path_name, doc_file)
            print("doc_path: ", doc_path)
            document = Document(doc_path)
            tables = document.tables
            print("tables: ", tables)
            print("len(tables): ", len(tables))
            doc_table_nums +=  len(tables)
        print("doc_table_nums: ", doc_table_nums)

        workbook = xlwt.Workbook(encoding='ascii')
        worksheet = workbook.add_sheet("MySheet")
        worksheet.write(0, 0, "监测编号")
        worksheet.write(0, 1, "变化后权属单位（个人）名称")
        worksheet.write(0, 2, "农用地转用、征地和供地批文号")
        worksheet.write(0, 3, "用途说明（建设动工、竣工时间及现状）")
        for doc_file in doc_files:
            # 获取文件的当前路径
            path = os.path.join(self.path_name, doc_file)
            print("path" + path)
            document = Document(path)
            tables = document.tables
            print(type(tables))
            print(len(tables))
            print("doc_file: ", doc_file)
            for i in range(doc_table_nums):
                table = tables[i]
                num = table.cell(0, 1).text
                print("num: ", num)
                person = table.cell(1, 1).text
                print("person: ", person)
                document = table.cell(2, 1).text
                print("document: ", document)
                explain = table.cell(3, 1).text
                print("explain: ", explain)
                print("I: ", i)
                worksheet.write(i+1, 0, num)
                worksheet.write(i+1, 1, person)
                worksheet.write(i+1, 2, document)
                worksheet.write(i+1, 3, explain)
            print("读取文件 %s 读取完毕。。。：" % doc_file)
        workbook.save('./extract_out/20223.xlsx')
        print("文件提取完毕，请查看extract文件夹下的xls文件")


if __name__ == '__main__':
    docs_path_name = r'./docs/'
    eui = ExtractNumsInfo(docs_path_name)
    eui.extract_info()
    print("不要急，饮茶先，马上就好...")
    time.sleep(3)
    eui.merge_excel_sheet()